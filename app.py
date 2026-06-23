# -*- coding: utf-8 -*-
"""
Formulario de verificación de datos visita - Optimizado para Vista Móvil
Se fuerza un layout responsivo tipo App Móvil/Smartphone y mantiene la lógica original de 'MUESTRA_FINAL'.
"""

import io
import json
from datetime import datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from streamlit_js_eval import get_geolocation
    GEO_OK = True
except Exception:
    GEO_OK = False

# --------------------------------------------------------------------------
# CONFIGURACIÓN GENERAL (Forzar contenedor móvil por CSS)
# --------------------------------------------------------------------------
st.set_page_config(
    page_title="App Visitas",
    page_icon="🏦",
    layout="centered",  # Centrado ayuda a emular una interfaz móvil en pantallas grandes
    initial_sidebar_state="collapsed",
)

EXCEL_COLUMNS = [
    "RECNO", "PEPAIS", "PETDOC", "PENDOC", "CODCLI", "BCEMP", "BCSUC", "BCMDA",
    "BCPAP", "BCCTA", "BCOPER", "BCSBOP", "BCTOP", "BCMOD", "CODCRE", "REGION",
    "ZONA", "AGENCIA", "CLIENTE", "DIRECCION_DOM", "DISTRITO_DOM",
    "PROVINCIA_DOM", "DEPARTAMENTO_DOM", "DIRECCION_NEG", "DISTRITO_NEG",
    "PROVINCIA_NEG", "DEPARTAMENTO_NEG", "ACTIVIDAD_ECON", "ANALISTA",
    "PRODUCTO_CAJA", "SALDO_MN", "SALDO_VIGE", "SALDO_REFI", "SALDO_VENC",
    "SALDO_JUDI", "MORA_CONT", "TIPO_SBS", "FECDES", "IMPDESEMB_MN",
    "COD_MODULO", "MODULO", "COD_TIPO_OPERACION", "TIPO_OPERACION",
    "ANALISTA_EVAL", "USUARIO_APROB", "USUARIO_DESEM", "FECHA_EVAL",
    "DIAS_ATRASO", "ESTADO_CREDITO", "ATRANT_1M", "ATRANT_2M", "ATRANT_3M",
    "ATRANT_4M", "ATRANT_5M", "ATRANT_6M", "TIPO_SOLI", "NUMERO_CUOTAS",
    "CUOTAS_PAGADAS", "TIPO", "SEGMENTACION_MYPE", "CATEG_RESULTANTE",
    "CATEG_RESULTANTE_SINALIN", "CUENTA_AVAL", "FECHA_UTLPAGO", "UAI_IND",
    "ESTRATO", "TIPO_EXPEDIENTE",
]

NARANJA = "C8102E"   
AZUL = "1B3A5C"
VERDE = "137333"
ROJO = "a50e0e"

# CSS Avanzado para forzar comportamiento de Smartphone App (Ancho máximo, padding y fuentes)
MOBILE_CSS = f"""
<style>
/* Forzar contenedor central simulando tamaño de dispositivo móvil */
.block-container {{
    max-width: 480px !important;
    padding-top: 1rem !important;
    padding-bottom: 5rem !important;
    padding-left: 12px !important;
    padding-right: 12px !important;
}}

.stApp {{ background-color: #F8F9FA; }}

/* Tarjetas/Cards Móviles */
.mockup-card {{
    background: white; 
    padding: 1.1rem; 
    border-radius: 14px;
    box-shadow: 0 4px 12px rgba(0,0,0,0.05); 
    margin-bottom: 1rem;
    border: 1px solid #E9ECEF;
}}

/* Encabezado fijo de perfil de usuario en la app */
.client-profile-box {{
    background-color: #E6F4EA; 
    border-radius: 12px; 
    padding: 0.8rem 1rem;
    margin-bottom: 1rem; 
    border: 1px solid #D1E7DD;
    font-size: 0.9rem;
}}

/* Botones Grandes estilo Mobile-First (Fácil de presionar con el pulgar) */
div.stButton > button {{
    background-color: #0052CC; 
    color: white; 
    border: none;
    border-radius: 10px; 
    padding: 12px 20px; 
    font-weight: 600; 
    font-size: 1rem;
    width: 100% !important;
    display: block;
    margin-bottom: 0.5rem;
}}
div.stButton > button:hover {{ background-color: #0043A4; color: white; }}

/* Estilo para botón de regreso */
div.stButton > button[key^="btn_atras"] {{
    background-color: #FFFFFF !important; 
    color: #{AZUL} !important;
    border: 1px solid #CED4DA !important;
}}

/* Panel de Validación de Riesgos */
.validation-box {{
    border: 1.5px solid #{AZUL}; 
    border-radius: 12px; 
    padding: 1rem;
    background: white; 
    margin: 0.8rem 0;
}}
.validation-title {{
    font-size: 1rem; 
    font-weight: 700; 
    color: #{AZUL}; 
    margin-bottom: 0.5rem;
    border-bottom: 2px solid #{NARANJA}; 
    padding-bottom: 0.2rem;
}}
.validation-box div.stButton > button {{
    background-color: #F8F9FA !important; 
    color: #1B3A5C !important;
    border: 1px solid #DEE2E6 !important; 
    text-align: left !important;
    padding: 10px !important; 
    margin-bottom: 6px !important;
    font-size: 0.85rem !important;
    border-radius: 8px !important;
}}

.badge-ok {{ background:#e6f4ea; color:#137333; padding:5px 10px; border-radius:10px; font-size:0.8rem; font-weight:600; display:inline-block; }}
.badge-pend {{ background:#fce8e6; color:#a50e0e; padding:5px 10px; border-radius:10px; font-size:0.8rem; font-weight:600; display:inline-block; }}

/* Ajuste de inputs para pantallas táctiles */
.stTextInput>div>div>input, .stSelectbox>div>div>div {{
    padding: 10px !important;
    font-size: 0.95rem !important;
}}
</style>
"""
st.markdown(MOBILE_CSS, unsafe_allow_html=True)

# --------------------------------------------------------------------------
# HELPERS DE INICIALIZACIÓN
# --------------------------------------------------------------------------
def safe_str(v, default=""):
    if v is None: return default
    try:
        if pd.isna(v): return default
    except Exception: pass
    s = str(v).strip()
    return default if s.lower() in ("nan", "none") else s

def safe_float(v, default=0.0):
    try:
        f = float(v)
        if pd.isna(f): return default
        return f
    except Exception: return default

def fmt_money(v):
    return f"S/. {safe_float(v):,.2f}"

def limpiar_texto_dni(val):
    if pd.isna(val) or val is None: return ""
    txt = str(val).strip()
    if txt.endswith(".0"): txt = txt[:-2]
    if txt.isdigit() and len(txt) < 8 and len(txt) > 0: txt = txt.zfill(8)
    return txt

if "step" not in st.session_state: st.session_state.step = "Búsqueda y Carga"
if "clientes_df" not in st.session_state: st.session_state.clientes_df = None
if "cliente_actual" not in st.session_state: st.session_state.cliente_actual = {}
if "visitas" not in st.session_state: st.session_state.visitas = {}
if "garantias" not in st.session_state: st.session_state.garantias = []
if "rcc" not in st.session_state: st.session_state.rcc = []
if "validaciones_marcadas" not in st.session_state: st.session_state.validaciones_marcadas = {}
if "click_timestamps" not in st.session_state: st.session_state.click_timestamps = {}

cliente = st.session_state.cliente_actual

# --------------------------------------------------------------------------
# VALIDACIONES AUTOMÁTICAS
# --------------------------------------------------------------------------
def validar_visita():
    validaciones = {
        "documentos_enmiendas": False, "documentos_inconsistentes": False,
        "documentos_sin_datos": False, "documentos_sin_firmas": False,
        "documentos_duplicados": False, "sin_sustento_actividad": False,
        "sin_sustento_ingresos": False, "sin_sustento_activos": False,
        "conyuge_omitido": False, "credito_reprogramado": False,
        "credito_refinanciado": False, "calificacion_diferente": False,
    }
    if not safe_str(cliente.get("CLIENTE")):
        validaciones["documentos_sin_datos"] = True
    if safe_str(cliente.get("DIAS_ATRASO")) and int(safe_float(cliente.get("DIAS_ATRASO"))) > 0:
        validaciones["calificacion_diferente"] = True
    visitas = st.session_state.visitas
    for clave in ["domicilio", "negocio", "aval"]:
        if clave not in visitas:
            validaciones["sin_sustento_actividad"] = True
            break
        if not visitas[clave].get("foto_bytes"):
            validaciones["documentos_sin_firmas"] = True
    return validaciones

def mostrar_panel_validacion():
    st.markdown('<div class="validation-box">', unsafe_allow_html=True)
    st.markdown('<div class="validation-title">🔍 Criterios de Riesgo (Touch)</div>', unsafe_allow_html=True)
    
    validaciones_auto = validar_visita()
    criterios = {
        "documentos_enmiendas": ("Docs con enmiendas", "⚠️"),
        "documentos_inconsistentes": ("Datos inconsistentes", "⚠️"),
        "documentos_sin_datos": ("Docs sin datos de cliente", "❌"),
        "documentos_sin_firmas": ("Docs sin firmas o fotos", "❌"),
        "documentos_duplicados": ("Documentos duplicados", "⚠️"),
        "sin_sustento_actividad": ("Sin sustento actividad", "❌"),
        "sin_sustento_ingresos": ("Sin sustento ingresos", "❌"),
        "sin_sustento_activos": ("Sin sustento activos", "⚠️"),
        "conyuge_omitido": ("Cónyuge omitido", "⚠️"),
        "credito_reprogramado": ("Crédito reprogramado", "ℹ️"),
        "credito_refinanciado": ("Crédito refinanciado", "ℹ️"),
        "calificacion_diferente": ("Calificación desactualizada", "⚠️"),
    }
    
    items_por_categoria = {"❌": [], "⚠️": [], "ℹ️": []}
    for key, (label, icon) in criterios.items():
        items_por_categoria[icon].append((key, label))
    
    with st.container(height=220, border=True):
        for icon in ["❌", "⚠️", "ℹ️"]:
            for key, label in items_por_categoria[icon]:
                is_checked = st.session_state.validaciones_marcadas.get(key, validaciones_auto.get(key, False))
                marcador_visual = "🔘" if is_checked else "⚪"
                
                if st.button(f"{marcador_visual} {icon} {label}", key=f"btn_crit_{key}", use_container_width=True):
                    ahora = datetime.now().timestamp()
                    ultimo_clic = st.session_state.click_timestamps.get(key, 0)
                    st.session_state.click_timestamps[key] = ahora
                    
                    if not is_checked:
                        st.session_state.validaciones_marcadas[key] = True
                        st.rerun()
                    else:
                        if (ahora - ultimo_clic) < 0.8: # Doble clic táctil
                            st.session_state.validaciones_marcadas[key] = False
                            st.rerun()
                            
    total_m = sum(1 for v in st.session_state.validaciones_marcadas.values() if v)
    if total_m == 0: st.success("✅ Sin riesgos marcados")
    else: st.warning(f"⚠️ {total_m} marcados | Doble toque para limpiar")
    st.markdown('</div>', unsafe_allow_html=True)

# --------------------------------------------------------------------------
# PERFIL FLOTANTE SUPERIOR MÓVIL
# --------------------------------------------------------------------------
if cliente and st.session_state.step != "Búsqueda y Carga":
    mora = safe_str(cliente.get("DIAS_ATRASO"), "0")
    color_mora = ROJO if int(safe_float(mora)) > 0 else VERDE
    st.markdown(f"""
    <div class="client-profile-box">
        <b>👤 {safe_str(cliente.get('CLIENTE'))}</b><br>
        <span style="color:#555;">DNI: {safe_str(cliente.get('PENDOC'))}</span> | 
        <span style="color:#{color_mora}; font-weight:bold;">Mora: {mora} días</span>
    </div>
    """, unsafe_allow_html=True)


# ==========================================================================
# 1️⃣ PASO: BÚSQUEDA Y CARGA (Para pantallas verticales)
# ==========================================================================
if st.session_state.step == "Búsqueda y Carga":
    st.markdown("### 🏦 Carga e Ingreso Móvil")
    
    st.markdown('<div class="mockup-card">', unsafe_allow_html=True)
    st.subheader("📂 Base de Datos de la Cartera")
    filas_a_saltar = st.number_input("Saltar filas cabecera:", min_value=0, value=0)
    excel_file = st.file_uploader("Subir archivo (.xlsx)", type=["xlsx"])
    
    if excel_file is not None:
        try:
            excel_lector = pd.ExcelFile(excel_file)
            hoja_objetivo = "MUESTRA_FINAL"
            if hoja_objetivo not in excel_lector.sheet_names:
                hoja_objetivo = excel_lector.sheet_names[0]
                st.warning(f"Usando primera pestaña: '{hoja_objetivo}'")
            
            df_cargado = pd.read_excel(excel_file, sheet_name=hoja_objetivo, skiprows=filas_a_saltar, dtype=str)
            df_cargado.columns = [str(c).strip().upper() for c in df_cargado.columns]
            
            if len(df_cargado.columns) >= 4:
                df_cargado = df_cargado.rename(columns={df_cargado.columns[3]: "PENDOC"})
            if "PENDOC" in df_cargado.columns:
                df_cargado["PENDOC"] = df_cargado["PENDOC"].apply(limpiar_texto_dni)
                
            st.session_state.clientes_df = df_cargado
            st.success(f"📊 {len(df_cargado)} filas leídas de '{hoja_objetivo}'")
        except Exception as e:
            st.error(f"Error: {e}")
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="mockup-card">', unsafe_allow_html=True)
    st.subheader("🔍 Localizador de Titular")
    busq = st.text_input("Ingresa DNI, Nombre o Código:")
    
    df = st.session_state.clientes_df
    if df is not None and busq:
        b = busq.strip().lower()
        mask = (
            df.get("PENDOC", pd.Series("", index=df.index)).astype(str).str.contains(b, case=False, na=False) |
            df.get("CODCLI", pd.Series("", index=df.index)).astype(str).str.contains(b, case=False, na=False) |
            df.get("CLIENTE", pd.Series("", index=df.index)).astype(str).str.contains(b, case=False, na=False)
        )
        resultados = df[mask]
        
        if len(resultados) > 0:
            opciones = resultados.apply(lambda r: f"{safe_str(r.get('CODCLI'))} - {safe_str(r.get('CLIENTE'))}", axis=1).tolist()
            sel = st.selectbox("Coincidencias halladas:", opciones)
            if sel:
                idx_sel = opciones.index(sel)
                if st.button("🔴 Cargar Ficha Cliente"):
                    st.session_state.cliente_actual = resultados.iloc[idx_sel].to_dict()
                    st.session_state.visitas = {}
                    st.session_state.validaciones_marcadas = {}
                    st.session_state.click_timestamps = {}
                    st.session_state.garantias = []
                    st.session_state.rcc = []
                    st.session_state.step = "Ficha del Cliente"
                    st.rerun()
        else:
            st.warning("Sin registros en MUESTRA_FINAL.")
    st.markdown('</div>', unsafe_allow_html=True)


# ==========================================================================
# 2️⃣ PASO: FICHA DEL CLIENTE (Estructura apilada)
# ==========================================================================
elif st.session_state.step == "Ficha del Cliente":
    st.markdown("### 💳 Ficha y Estado")
    
    st.markdown('<div class="mockup-card">', unsafe_allow_html=True)
    st.write(f"**Desembolsado:** {fmt_money(cliente.get('IMPDESEMB_MN'))}")
    st.write(f"**Saldo Capital:** {fmt_money(cliente.get('SALDO_MN'))}")
    st.write(f"**Calificación:** {safe_str(cliente.get('CATEG_RESULTANTE', '-'))}")
    st.markdown('</div>', unsafe_allow_html=True)

    with st.expander("📝 Detalles Crediticios Completos", expanded=False):
        st.text_input("Agencia", value=safe_str(cliente.get("AGENCIA")), disabled=True)
        st.text_input("Código", value=safe_str(cliente.get("CODCLI")), disabled=True)
        st.text_input("Operación", value=safe_str(cliente.get("BCOPER")), disabled=True)
        st.text_input("Producto", value=safe_str(cliente.get("PRODUCTO_CAJA")), disabled=True)
        st.text_input("Actividad", value=safe_str(cliente.get("ACTIVIDAD_ECON")), disabled=True)

    mostrar_panel_validacion()

    # Botones apilados verticalmente para facilitar pulsación con una mano
    if st.button("Continuar a Visitas de Campo ➡️"):
        st.session_state.step = "Visita"
        st.rerun()
    if st.button("⬅️ Volver a Carga", key="btn_atras_f"):
        st.session_state.step = "Búsqueda y Carga"
        st.rerun()


# ==========================================================================
# 3️⃣ PASO: VISITA (Optimizado para pantalla de celular y captura de cámara)
# ==========================================================================
elif st.session_state.step == "Visita":
    st.markdown("### 📍 Información de Campo")
    
    punto_v = st.radio("Entorno verificado:", ["Domicilio", "Negocio", "Aval"], horizontal=True)
    clave_v = punto_v.lower()
    
    st.markdown('<div class="mockup-card">', unsafe_allow_html=True)
    if clave_v == "domicilio":
        st.text_input("Dirección", value=safe_str(cliente.get("DIRECCION_DOM")))
        st.text_input("Distrito", value=safe_str(cliente.get("DISTRITO_DOM")))
        st.text_area("Referencia de Acceso", key="ref_dom")
        st.selectbox("Estructura Patrimonial", ["Propia", "Familiar", "Alquilada", "Otro"])
    elif clave_v == "negocio":
        st.text_input("Dirección Comercial", value=safe_str(cliente.get("DIRECCION_NEG")))
        st.text_input("Distrito Comercial", value=safe_str(cliente.get("DISTRITO_NEG")))
        st.text_area("Referencia Comercial", key="ref_neg")
        st.text_input("Giro Detallado", value=safe_str(cliente.get("ACTIVIDAD_ECON")))
    else:
        st.text_input("Código / Cuenta del Aval", value=safe_str(cliente.get("CUENTA_AVAL")))
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="mockup-card">', unsafe_allow_html=True)
    st.subheader("Geolocalización y Foto")
    
    visitas_data = st.session_state.visitas.get(clave_v, {})
    lat, lon = visitas_data.get("lat"), visitas_data.get("lon")
    
    entrevista_con = st.text_input("Entrevistado:", key=f"e_{clave_v}")
    
    if st.button("📡 Capturar Coordenadas GPS", key=f"gps_btn_{clave_v}"):
        if GEO_OK:
            loc = get_geolocation(key=f"geo_{clave_v}_{datetime.now().timestamp()}")
            if loc and "coords" in loc:
                lat, lon = loc["coords"]["latitude"], loc["coords"]["longitude"]
        else:
            st.warning("GPS no disponible.")
            
    if lat and lon:
        st.success(f"Ubicación: {lat:.5f}, {lon:.5f}")
        st.map(pd.DataFrame({"lat": [lat], "lon": [lon]}), zoom=15, height=180)

    st.markdown("---")
    f_cam = st.camera_input("Foto instantánea de campo", key=f"cam_{clave_v}")
    f_gal = st.file_uploader("O adjuntar de la galería", type=["png", "jpg", "jpeg"], key=f"gal_{clave_v}")
    foto_final = f_cam if f_cam is not None else f_gal

    comentarios = st.text_area("Observaciones:", key=f"com_{clave_v}")

    if st.button(f"💾 Guardar Visita {punto_v}"):
        st.session_state.visitas[clave_v] = {
            "fecha": str(datetime.now().date()), "hora": str(datetime.now().time())[:5], "entrevista_con": entrevista_con,
            "comentarios": comentarios, "lat": lat, "lon": lon,
            "foto_bytes": foto_final.getvalue() if foto_final is not None else None
        }
        st.success(f"✅ Datos de {punto_v} guardados.")
    st.markdown('</div>', unsafe_allow_html=True)

    if st.button("Ir a Evaluación Financiera ➡️"):
        st.session_state.step = "Ingresos y Gastos"
        st.rerun()
    if st.button("⬅️ Volver a Ficha", key="btn_atras_v"):
        st.session_state.step = "Ficha del Cliente"
        st.rerun()


# ==========================================================================
# 4️⃣ PASO: INGRESOS Y GASTOS (Vista de Inputs continuos verticales)
# ==========================================================================
elif st.session_state.step == "Ingresos y Gastos":
    st.markdown("### 📊 Datos Financieros")
    
    st.markdown('<div class="mockup-card">', unsafe_allow_html=True)
    ventas = st.number_input("Ventas Totales Mensuales (S/.)", value=0.0)
    costo_ventas = st.number_input("Costo de Ventas / Insumos (S/.)", value=0.0)
    gastos_admin = st.number_input("Gastos del Local / Alquiler (S/.)", value=0.0)
    gastos_financieros = st.number_input("Gastos Financieros (S/.)", value=0.0)
    otros_ingresos = st.number_input("Otros Ingresos (S/.)", value=0.0)
    gastos_familiares = st.number_input("Canasta Familiar (S/.)", value=0.0)
    st.markdown('</div>', unsafe_allow_html=True)

    resultado_neto = ventas + otros_ingresos - costo_ventas - gastos_admin - gastos_financieros - gastos_familiares
    utilidad_neta = resultado_neto - gastos_familiares

    st.markdown(f"""
    <div style="background-color: white; border-radius:12px; padding: 1rem; border: 1px solid #D2D7DF; text-align:center; margin-bottom: 1rem;">
        <small style="color:#666;">UTILIDAD NETA MENSUAL</small>
        <h3 style="color:#{VERDE}; margin: 2px 0 0 0; font-size:1.6rem;">{fmt_money(utilidad_neta)}</h3>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Ir al Reporte y Cierre ➡️"):
        st.session_state.step = "Reporte"
        st.rerun()
    if st.button("⬅️ Volver a Visitas", key="btn_atras_i"):
        st.session_state.step = "Visita"
        st.rerun()


# ==========================================================================
# 5️⃣ PASO: REPORTE (Cierre operativo y Descarga móvil)
# ==========================================================================
elif st.session_state.step == "Reporte":
    st.markdown("### 🏁 Descarga de Expediente")
    
    st.markdown('<div class="mockup-card">', unsafe_allow_html=True)
    status_dom = "🟢 Dom. Ok" if "domicilio" in st.session_state.visitas else "🔴 Dom. Pendiente"
    status_neg = "🟢 Neg. Ok" if "negocio" in st.session_state.visitas else "🔴 Neg. Pendiente"
    
    st.write(f"• **Estado Domicilio:** {status_dom}")
    st.write(f"• **Estado Negocio:** {status_neg}")
    st.markdown('</div>', unsafe_allow_html=True)

    # REUTILIZACIÓN COMPLETA DEL MOTOR DE EXPORTACIÓN ORIGINAL (.DOCX)
    def add_heading(doc, text, size=13, color=AZUL):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        return p

    def add_kv_table(doc, pairs, cols=2):
        table = doc.add_table(rows=0, cols=cols * 2)
        table.style = "Light Grid Accent 1"
        row = None
        for i, (k, v) in enumerate(pairs):
            if i % cols == 0: row = table.add_row().cells
            c = (i % cols) * 2
            row[c].text = str(k)
            row[c + 1].text = str(v) if v not in (None, "") else "-"
        return table

    def empaquetar_archivo_oficial():
        doc = Document()
        doc.add_heading("VISITA A CLIENTES DE PEQUEÑA EMPRESA", level=0)
        
        add_heading(doc, "I. Datos del cliente")
        add_kv_table(doc, [
            ("Titular", safe_str(cliente.get("CLIENTE"))), ("DNI", safe_str(cliente.get("PENDOC"))),
            ("Saldo capital", fmt_money(cliente.get("SALDO_MN"))), ("Días atraso", safe_str(cliente.get("DIAS_ATRASO")))
        ])
        
        # Inserción de fotos guardadas por la cámara del celular
        for clv, tit in [("domicilio", "III. Visita al domicilio"), ("negocio", "IV. Visita al negocio")]:
            add_heading(doc, tit)
            if clv in st.session_state.visitas:
                d = st.session_state.visitas[clv]
                add_kv_table(doc, [("Comentarios", d.get("comentarios")), ("GPS", f"{d.get('lat')},{d.get('lon')}")])
                if d.get("foto_bytes"):
                    doc.add_picture(io.BytesIO(d["foto_bytes"]), width=Cm(7))
            else:
                doc.add_paragraph("No registrado.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    st.markdown('<div class="mockup-card" style="text-align:center;">', unsafe_allow_html=True)
    try:
        archivo_word = empaquetar_archivo_oficial()
        nombre_salida = f"Informe_{safe_str(cliente.get('PENDOC', 'cliente'))}.docx"
        
        st.download_button(
            label="📥 Descargar Reporte (.docx)",
            data=archivo_word,
            file_name=nombre_salida,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        st.error(f"Error al compilar informe: {e}")
    st.markdown('</div>', unsafe_allow_html=True)

    if st.button("🔄 Evaluar Otro Cliente"):
        st.session_state.step = "Búsqueda y Carga"
        st.session_state.cliente_actual = {}
        st.session_state.visitas = {}
        st.session_state.garantias = []
        st.session_state.rcc = []
        st.session_state.validaciones_marcadas = {}
        st.session_state.click_timestamps = {}
        st.rerun()
